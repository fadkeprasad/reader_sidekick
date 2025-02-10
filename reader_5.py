import fitz  # PyMuPDF
import pyperclip
import os
import time
from win32gui import GetForegroundWindow, GetWindowText
from docx import Document

# Path to save the Word document (in the user's Documents folder)
DOCUMENT_PATH = r"C:\Users\fadke\Documents\Saved Quotes.docx"

def ensure_document_exists():
    """Ensure that the Word document exists and is accessible."""
    if not os.path.exists(DOCUMENT_PATH):
        doc = Document()
        doc.add_heading("Saved Quotes", level=1)
        doc.save(DOCUMENT_PATH)
        print(f"New document created at {DOCUMENT_PATH}")

def get_active_window_title():
    """Get the title of the currently active window."""
    return GetWindowText(GetForegroundWindow())

def extract_pdf_name(window_title):
    """Extract only the PDF name from the active window title."""
    if ".pdf" in window_title.lower():
        # Extract substring ending with '.pdf'
        start_index = window_title.lower().find(".pdf")
        pdf_name = window_title[:start_index + 4]  # Include '.pdf'
        return pdf_name.strip()
    return None

def find_pdf_path(pdf_name):
    """Search for the PDF file dynamically in common directories."""
    # Directories to search
    search_dirs = [
        os.path.expanduser("~/Documents"),  # User's Documents folder
        os.path.expanduser("~/Downloads"),  # User's Downloads folder
        os.path.expanduser("~/Desktop"),    # User's Desktop folder
    ]

    for search_dir in search_dirs:
        print(f"Searching for '{pdf_name}' in {search_dir} and subdirectories...")
        for root, _, files in os.walk(search_dir):
            for file in files:
                if file.lower() == pdf_name.lower():  # Case-insensitive match
                    return os.path.join(root, file)
    return None

def find_page_number(pdf_path, text):
    """Find the page number of a specific text in a PDF."""
    try:
        with fitz.open(pdf_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                if text in page.get_text():
                    return page_num + 1  # Page numbers are 1-based
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return "Unknown"

def append_to_word_doc(text, pdf_name, page_number):
    """Append text, PDF name, and page number to the Word document."""
    try:
        doc = Document(DOCUMENT_PATH)
        doc.add_paragraph(f"Quote: {text}")
        doc.add_paragraph(f"Source: {pdf_name}, Page: {page_number}")
        doc.add_paragraph("-" * 40)
        doc.save(DOCUMENT_PATH)
        print(f"Quote saved to {DOCUMENT_PATH}")
    except Exception as e:
        print(f"Error saving to document: {e}")

def monitor_clipboard():
    """Monitor clipboard for copied text."""
    print("Monitoring clipboard... Press 'Ctrl+C' in your PDF reader to copy text.")
    last_clipboard_content = ""

    while True:
        try:
            clipboard_content = pyperclip.paste()
            if clipboard_content != last_clipboard_content and clipboard_content.strip():
                last_clipboard_content = clipboard_content
                active_window_title = get_active_window_title()
                print(f"Active Window Title: {active_window_title}")

                pdf_name = extract_pdf_name(active_window_title)
                if pdf_name:
                    print(f"Detected PDF Name: {pdf_name}")
                    pdf_path = find_pdf_path(pdf_name)
                    if pdf_path:
                        print(f"PDF Path Found: {pdf_path}")
                        page_number = find_page_number(pdf_path, clipboard_content)
                        append_to_word_doc(clipboard_content, pdf_name, page_number)
                    else:
                        print(f"PDF file not found: {pdf_name}")
                        print("Suggestions:")
                        print("- Ensure the PDF is located in one of the searched directories.")
                        print("- Check for filename mismatches (e.g., spaces, capitalization).")
                        print("- Verify that the extracted name matches an actual file.")
                else:
                    print("No valid PDF detected.")
            time.sleep(0.5)

        except KeyboardInterrupt:
            print("\nExiting...")
            break
        except Exception as e:
            print(f"Unexpected error: {e}")
            time.sleep(1)

if __name__ == "__main__":
    ensure_document_exists()
    monitor_clipboard()